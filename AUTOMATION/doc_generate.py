from docx import Document
from docx.shared import Pt

# Create a Word document
doc = Document()
doc.add_heading('DAFTAR HADIR', level=1)
doc.add_paragraph('EMPLOYEE GATHERING IT APP DEV 2025')

# Create a table with 6 columns: No, NAMA, NIP, JABATAN, TANDA TANGAN (split into Left and Right)
table = doc.add_table(rows=48, cols=6)
table.style = 'Table Grid'

# Add header row
headers = ['No', 'NAMA', 'NIP', 'JABATAN', 'TANDA TANGAN', '']
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

# Merge header cells for "TANDA TANGAN"
table.cell(0, 4).merge(table.cell(0, 5))
table.cell(0, 4).text = 'TANDA TANGAN'

# Populate rows for 94 people (2 signatures per row in the "TANDA TANGAN" column)
person_num = 1
for row_idx in range(1, 48):  # 47 rows for attendees (94 people total)
    # Fill in row data
    table.cell(row_idx, 0).text = f"{person_num}\n{person_num + 1}"  # No for two people
    table.cell(row_idx, 1).text = ''  # NAMA
    table.cell(row_idx, 2).text = ''  # NIP
    table.cell(row_idx, 3).text = ''  # JABATAN
    
    # TANDA TANGAN for two people (left and right)
    table.cell(row_idx, 4).text = str(person_num)  # Left signature
    table.cell(row_idx, 5).text = str(person_num + 1)  # Right signature
    person_num += 2

# Adjust font size for all cells
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Save the document
file_path = "daftar_hadir_format.docx"
doc.save(file_path)
file_path